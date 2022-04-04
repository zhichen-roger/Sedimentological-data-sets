__author__ = "Huzhichen"
# coding=UTF-8
import os
from time import sleep

import nltk
from stanfordcorenlp import StanfordCoreNLP
import pyner
from nltk.corpus import brown
import re
from collections import Counter
from docx import Document
from docx.shared import RGBColor
brown_train = brown.tagged_sents(categories='news')
regexp_tagger = nltk.RegexpTagger(
    [(r'^-?[0-9]+(.[0-9]+)?$', 'CD'),
     (r'(-|:|;)$', ':'),
     (r'\'*$', 'MD'),
     (r'(The|the|A|a|An|an)$', 'AT'),
     #(r'.*able$', 'JJ'),
     (r'^±', 'JJ'),
     (r'^[A-Z].*$', 'NNP'),
     (r'.*ness$', 'NN'),
     (r'.*ly$', 'RB'),
     (r'.*s$', 'NNS'),
     (r'^~?[0-9]','VB'),
     (r'.*ing$', 'VBG'),
     (r'.*ed$', 'VBD'),
     #(r'et$','FW-CC'),
     #(r'al$','FW-NN'),
     (r'and$','CC'),
     (r'.*', 'NN')
     ])
unigram_tagger = nltk.UnigramTagger(brown_train, backoff=regexp_tagger)
bigram_tagger = nltk.BigramTagger(brown_train, backoff=unigram_tagger)

# This is our semi-CFG; Extend it according to your own needs
#############################################################################
cfg = {}
cfg["NNP+NNP"] = "NNP"
cfg["CD+NNP"] = "NNP"
cfg["VB+NNP"] = "NNP"
cfg["NN+NN"] = "NNI"
cfg["NNI+NN"] = "NNI"
cfg["CD+JJ"] = "NNP"
cfg["CD+NN"] = "NNP"
cfg["CD+IN"] = "NNP"
cfg["JJ+CD"] = "NNP"
cfg["JJ+JJ"] = "JJ"
cfg["JJ+NN"] = "NNI"
#cfg["FW-CC+FW-NN"] = "NNI"
cfg["CD+CC"] = "NNP"
#############################################################################
class NPExtractor(object):
    def __init__(self, sentence):
        self.sentence = sentence
    def dict(self):
        Dict = []
        # 打开字典并去掉换行符
        with open("../duc.txt", "r", encoding="utf-8") as f:
            lines = f.readlines()
            # 去除换行符
            result = ([x.strip() for x in lines if x.strip() != ''])
            # 将整理好字典提取成为全局字典
            for x in result:
                Dict.append(x)
        return Dict

    def geosubstance(self):
        geo = []
        # 打开字典并去掉换行符
        with open("../geosubstance.txt", "r", encoding="utf-8") as f:
            lines = f.readlines()
            # 去除换行符
            result = ([x.strip() for x in lines if x.strip() != ''])
            # 将整理好字典提取成为全局字典
            for x in result:
                geo.append(x)
        return geo
    # Split the sentence into singlw words/tokens
    def tokenize_sentence(self, sentence):
        tokens = nltk.word_tokenize(sentence)
        return tokens
    # Normalize brown corpus' tags ("NN", "NN-PL", "NNS" > "NN")
    def normalize_tags(self, tagged):
        n_tagged = []
        for t in tagged:
            if t[1] == "NP-TL" or t[1] == "NP":
                n_tagged.append((t[0], "NNP"))
                continue
            if t[1].endswith("-TL"):
                n_tagged.append((t[0], t[1][:-3]))
                continue
            if t[1].endswith("S"):
                n_tagged.append((t[0], t[1][:-1]))
                continue
            n_tagged.append((t[0], t[1]))
        return n_tagged
    # Extract the main topics from the sentence
    def extract(self):
        tokens = self.tokenize_sentence(self.sentence)
        tags = self.normalize_tags(bigram_tagger.tag(tokens))
        # print(tags)
        merge = True
        while merge:
            merge = False
            for x in range(0, len(tags) - 1):
                t1 = tags[x]
                t2 = tags[x + 1]
                key = "%s+%s" % (t1[1], t2[1])
                value = cfg.get(key, '')
                if value:
                    merge = True
                    tags.pop(x)
                    tags.pop(x)
                    match = "%s %s" % (t1[0], t2[0])
                    pos = value
                    tags.insert(x, (match, pos))
                    break
        matches = []
        for t in tags:
            if t[1] == "NNP" or t[1] == "NNI" or t[1] =="NNS":
                # if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NN":
                matches.append(t[0])
        return matches
    def extractSubstance(self):
        geo = self.geosubstance()
        tokens = self.tokenize_sentence(self.sentence)
        tags = self.normalize_tags(bigram_tagger.tag(tokens))
        # print(tags)
        merge = True
        while merge:
            merge = False
            for x in range(0, len(tags) - 1):
                t1 = tags[x]
                t2 = tags[x + 1]
                key = "%s+%s" % (t1[1], t2[1])
                value = cfg.get(key, '')
                if value:
                    merge = True
                    tags.pop(x)
                    tags.pop(x)
                    match = "%s %s" % (t1[0], t2[0])
                    pos = value
                    tags.insert(x, (match, pos))
                    break
        matches = []
        for t in tags:
            if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NNS":
                # if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NN":
                matches.append(t[0])
        result = []

        for m in range(len(matches)):
            for d in range(len(geo)):
                res = matches[m].lower().find(geo[d].lower())
                if (res != -1):
                    result.append(matches[m])
            if ( re.findall(r'[0-9]+ +[m|km|kilometer|meter]', matches[m]) ):
                result.append(matches[m])

        correct = []
        for item in range(len(result)):
            if 'Ma' not in result[item]:
                correct.append(result[item])

        seen = set()
        Substance = []
        for item in correct:
            if item not in seen:
                seen.add(item)
                Substance.append(item)
        return Substance
    def extractTime(self):
        dict = self.dict()
        tokens = self.tokenize_sentence(self.sentence)
        tags = self.normalize_tags(bigram_tagger.tag(tokens))
        print(tags)
        merge = True
        while merge:
            merge = False
            for x in range(0, len(tags) - 1):
                t1 = tags[x]
                t2 = tags[x + 1]
                key = "%s+%s" % (t1[1], t2[1])
                value = cfg.get(key, '')
                if value:
                    merge = True
                    tags.pop(x)
                    tags.pop(x)
                    match = "%s %s" % (t1[0], t2[0])
                    pos = value
                    tags.insert(x, (match, pos))
                    break
        matches = []
        for t in tags:
            if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NNS":
                # if t[1] == "NNP" or t[1] == "NNI" or t[1] == "NN":
                matches.append(t[0])
        result = []
        for m in range(len(matches)):
            for d in range(len(dict)):
                res = matches[m].lower().find(dict[d].lower())
                if (res != -1):
                    result.append(matches[m])
            #if (re.findall(r'[0-9]', matches[m]) and not(re.findall(r'Fig+[0-9]', matches[m])) and not(re.findall(r'[0-9]+ +[m|km|kilometer|meter]',matches[m]))):
            #if (re.findall(r'[0-9]+ +[Ma]', matches[m]) or re.findall(r'[0-9]+ +[and]', matches[m]) or re.findall(r'[0-9]+ +[to]', matches[m])):
            if re.findall(r'[0-9]+ +[Ma]', matches[m]):
                result.append(matches[m])
        # correct = []
        # for item in range(len(result)-1):
        #     if ('to' in result[item]  or 'and' in result[item] ):
        #         if 'Ma' in result[item + 1]:
        #             correct.append(result[item])
        #             correct.append(result[item + 1])
        #             item += 1
        #
        #     elif 'Ma' in result[item]:
        #         correct.append(result[item])
        #     else:
        #         correct.append(result[item])
        seen = set()
        Time = []
        for item in result:
            if item not in seen:
                seen.add(item)
                Time.append(item)
        return Time
    def extractLocation(self,sentence):
        nlp = StanfordCoreNLP('E:\stanford\stanford-corenlp-latest\stanford-corenlp-full-2021-01-09', lang='en')
        Information = []
        Location = []
        for i in range(0, len(nlp.ner(sentence))):
            if nlp.ner(sentence)[i][1] == 'LOCATION':
                Information.append(nlp.ner(sentence)[i][0])
                #print(nlp.ner(sentence)[i][0])  # 实体识别
        seen = set()
        for item in Information:
            if item not in seen:
                seen.add(item)
                Location.append(item)
        nlp.close()
        return Location
# Main method"
# 时间存储词典
def dict():
    # 打开字典并去掉换行符
    with open("../duc.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
        # 去除换行符
        result = ([x.strip() for x in lines if x.strip() != ''])
        # 将整理好字典提取成为全局字典
        for x in result:
            DictTim.append(x)
    return DictTim
# 地点存储词典
def geodict():
    # 打开字典并去掉换行符
    with open("../geosubstance.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
        # 去除换行符
        result = ([x.strip() for x in lines if x.strip() != ''])
        # 将整理好字典提取成为全局字典
        for x in result:
            GeoDict.append(x)
    return GeoDict

def extractParagraph():
    # countFP = 0
    countTxt = 0
    osfile = []
    dir = r"F:\pythonbert\demo\untitled1\CarbonatePlatform\extracttext"
    for root, dirs, files in os.walk(dir):
        for file in files:
            if file != 'extract.py' and file != '__init__.py' and file !='split.py' and file !='Label.py' and file !='CSV.py' and file !='resplit.py' and file !='cluster.py':
                osfile.append(file)
    for index in range(len(osfile)):
        Geo_Tim = []
        countFP = 0
        # document = Document()
        # file = document.add_paragraph()
        # file1 = document.add_paragraph()
        re = open('../Txtextract/'+'extract_'+osfile[index].split('.')[0]+'.txt', "w", encoding='utf-8')
        Fp = open('../Txtextract/' + 'Formation_platform_' + osfile[index].split('.')[0] + '.txt', "w", encoding='utf-8')
        fpcount = open('../doc/results/result_' +osfile[index].split('.')[0]+ '.txt', "w", encoding='utf-8')
        with open(osfile[index], "r", encoding='utf-8') as f:
            while True:
                data = f.readline()
                np_extractor = NPExtractor(data)
                result = np_extractor.extract()
                substance = np_extractor.extractSubstance()
                time = np_extractor.extractTime()
                location = np_extractor.extractLocation(data)

                # Listsubstance == substance
                # ListTim == time

                # 进行每句切词分类
                geoList = []
                timList = []
                for i in range(len(substance)):
                    flag = 0
                    for g in range(len(GeoDict)):
                        if (substance[i].find(GeoDict[g]) != -1):
                            flag = 1
                    if flag == 1:
                        geoList.append(substance[i])
                for i in range(len(time)):
                    flag = 0
                    for g in range(len(DictTim)):
                        if (time[i].find(DictTim[g]) != -1):
                            flag = 1
                    if flag == 1:
                        timList.append(time[i])
                # 这里
                geo_tim = []
                mix = []
                geofinall =[]
                timfinall=[]
                for i in range(len(geoList)):
                    fig = 0
                    for g in range(len(DictTim)):
                        if (geoList[i].find(DictTim[g]) != -1):
                            fig = 1
                    if fig != 1:
                        geofinall.append(geoList[i].strip())
                    elif fig == 1:
                        mix.append(geoList[i].strip())

                for i in range(len(timList)):
                    fig = 0
                    for g in range(len(GeoDict)):
                        if (timList[i].find(GeoDict[g]) != -1):
                            fig = 1
                    if fig != 1:
                        timfinall.append(timList[i].strip())

                for index in range(len(mix)):
                    stra = mix[index]
                    strtmp = stra.split(' ')
                    max = 0
                    # print(strtmp)
                    for s in range(len(strtmp)):
                        for g in range(len(DictTim)):
                            if DictTim[g] in strtmp[s]:
                                if max < s:
                                    max = s
                    # print(max)
                    strtim = ''
                    for i in range(max + 1):
                        strtim += strtmp[i] + ' '
                    # print("zhes", strtim)
                    timfinall.append(strtim.strip())

                    strgeo = ''
                    for i in range(max + 1, len(strtmp)):
                        strgeo += strtmp[i] + ' '
                    # print("geozhes", strgeo)
                    geofinall.append(strgeo.strip())

                geo_tim = []

                for index in range(len(geofinall)):
                    for t in range(len(timfinall)):
                        strgt = geofinall[index]+" + "+timfinall[t]
                        geo_tim.append(strgt)
                        Geo_Tim.append(strgt)

                print("This sentence is about: %s" % ", ".join(result))
                print("This substance is about: %s" % ", ".join(geofinall))
                print("This time is about: %s" % ", ".join(timfinall))
                print("This Location is about: %s" % ", ".join(location))
                print("This mix is about: %s" % ", ".join(geo_tim))

                # 分类写入txt
                if not data:
                    break
                if not geofinall or not timfinall:
                    continue
                re.write(data)
                re.write("Key words: %s" % ", ".join(result)+'\n')
                re.write("Substance: %s" % ", ".join(geofinall)+'\n')
                re.write("Time: %s" % ", ".join(timfinall) + '\n')
                re.write("Location: %s" % ", ".join(location) + '\n')
                re.write("mix: %s" % ", ".join(geo_tim) + '\r\n')
                # 分类标红写入docx
                # 统计Formation_platform 写入txt
                flagF = 0
                flagP = 0
                for index in range(len(geofinall)):
                    if geofinall[index].find('platform') != -1:
                        flagF = 1
                    elif geofinall[index].find('Formation') != -1:
                        flagP = 1

                # if len(geo_tim) == 1:
                #     file.add_run(data)
                #     file.add_run("Key words: %s" % ", ".join(result)+'\n')
                #     file.add_run("Substance: %s" % ", ".join(geofinall)+'\n')
                #     file.add_run("Time: %s" % ", ".join(timfinall) + '\n')
                #     file.add_run("Location: %s" % ", ".join(location) + '\n')
                #     run = file.add_run("mix: %s" % ", ".join(geo_tim) + '\r\n')
                #     run.font.color.rgb = RGBColor(250,0,0)
                #     # 设置加粗
                #     run.font.bold = True
                # elif flagF == 1 and flagP == 1 :
                #     countTxt += 1
                #     run1 = file.add_run(data)
                #     run1.font.color.rgb = RGBColor(0, 0, 250)
                #     run1.font.bold = True
                #     file.add_run("Key words: %s" % ", ".join(result) + '\n')
                #     run2 = file.add_run("Substance: %s" % ", ".join(geofinall) + '\n')
                #     run2.font.color.rgb = RGBColor(0, 0, 250)
                #     run2.font.bold = True
                #     file.add_run("Time: %s" % ", ".join(timfinall) + '\n')
                #     file.add_run("Location: %s" % ", ".join(location) + '\n')
                #     file.add_run("mix: %s" % ", ".join(geo_tim) + '\r\n')
                # else:
                #     file.add_run(data)
                #     file.add_run("Key words: %s" % ", ".join(result) + '\n')
                #     file.add_run("Substance: %s" % ", ".join(geofinall) + '\n')
                #     file.add_run("Time: %s" % ", ".join(timfinall) + '\n')
                #     file.add_run("Location: %s" % ", ".join(location) + '\n')
                #     file.add_run("mix: %s" % ", ".join(geo_tim) + '\r\n')
                # file.add_run(str(countTxt))
                if flagF == 1 and flagP == 1:
                    countFP += 1
                    Fp.write(data)
                    Fp.write("Formation_platform: %s" % ", ".join(geofinall) + '\r\n')
        Fp.write('Count Formation_platform: '+str(countFP))
        # document.save('../Txtextract/'+'extract_'+osfile[index].split('.')[0]+'.docx')
        re.close()
        Fp.close()
        #sleep(100)
        georesult = Counter(Geo_Tim)
        g = georesult.most_common(len(georesult))
        strrgeo = []
        for i in g:
            strrgeo.append(str(i[1]) + " " + str(i[0]))
        fpcount.write("Count Substance + Time:\n")
        fpcount.write("%s" % '\n'.join(strrgeo) + '\n')
        fpcount.close()


# def counterTxt():
#     georesult = Counter(Geo_Tim)
#     g = georesult.most_common(len(georesult))
#     f = open('../doc/results/result_'+'.txt', "w", encoding='utf-8')
#     strrgeo = []
#     for i in g:
#         strrgeo.append(str(i[1]) + " " + str(i[0]))
#     f.write("Count Substance + Time:\n")
#     f.write("%s" % '\n'.join(strrgeo) + '\n')

if __name__ == '__main__':
    # 读取词典
    DictTim = []
    GeoDict = []
    dict()
    geodict()
    # 反抽主函数
    #Geo_Tim = []
    extractParagraph()
    #counterTxt()

