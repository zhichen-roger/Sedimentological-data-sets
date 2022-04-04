__author__ = "Huzhichen"
from docx import Document
import  re
import os
# 时间存储字典
def dict():
    # 打开字典并去掉换行符
    with open("../duc.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
        # 去除换行符
        result = ([x.strip() for x in lines if x.strip() != ''])
        # 将整理好字典提取成为全局字典
        for x in result:
            Dict.append(x)
    return Dict

# 地点存储词典
def geodict():
    # 打开字典并去掉换行符
    with open("../geosubstance.txt", "r", encoding="utf-8") as f:
        lines = f.readlines()
        # 去除换行符
        result = ([x.strip() for x in lines if x.strip() != ''])
        # 将整理好字典提取成为全局字典
        for x in result:
            GeoDict.append(x)
    return GeoDict

# 读取每段文本
# 进行初步清理并合并成一个文本
def getParagraphsText(doc,Namedocx):
    for i in range(0,len(doc.paragraphs)):
        print("这是:",doc.paragraphs[i].text)
    for i in range(0, len(doc.paragraphs)):
        paragraphsText.append(doc.paragraphs[i].text)
    text = ([x.strip() for x in paragraphsText if x.strip() != ''])
    # 初步数据清理从abstract开始删除acknowedge以后字段
    a = 0
    for i in range(len(text)):
        if text[i] == "abstract":
            a = i
    for i in range(a+1,len(text)):
        if text[i] == 'Acknowledgements':
            break
        else:
            paragraphsText_new.append(text[i])
    f = open('../txt/' + Namedocx.split('.')[0] + ".txt", "w", encoding='utf-8')
    for line in paragraphsText_new:
        f.write(' '+line)
    print(Namedocx + "保存成功")
    f.close()

# 进行提取筛选
def filter():
    osfile = []
    result = []
    splitTxt = []
    List = []
    ListDict = []
    dir = r"D:\pythonProject\CarbonatePlatform\txt"
    for root, dirs, files in os.walk(dir):
        for file in files:
            if file.split('.')[1] == 'txt':
                osfile.append(file)
    for index in range(len(osfile)):
        with open('../txt/' + osfile[index], "r", encoding="utf-8") as f:
            lines = f.readlines()
            result = ([x.strip() for x in lines if x.strip() != ''])
        for index in range(len(result)):
           splitTxt=result[index].replace("ca.","ca").replace("e.g.,","").replace("e.g.","").replace("al.","al")\
               .replace("Fig. 1","").replace("Fig. 2","").replace("Fig. 3","").replace("Fig. 4","").replace("Fig. 5","").replace("Fig. 6","").replace("Fig. 7","").replace("Fig. 8","").replace("Fig. 8","").replace("Fig. 9","").replace("Fig. 10","").replace("Fig. 11","")\
               .replace("Fig. 1.","").replace("Fig. 2.","").replace("Fig. 3.","").replace("Fig. 4.","").replace("Fig. 5.","").replace("Fig. 6.","").replace("Fig. 7.","").replace("Fig. 8.","").replace("Fig. 9.","").replace("Fig. 10.","").replace("Fig. 11.","")\
               .replace("cf.","cf").replace("Fig.","").split('.')
        flag1geo = 0
        flag1dict = 0
        flag2geo = 0
        flag2dict = 0
        for i in range(len(splitTxt)-1):
           if (flag1geo == 1 and flag2dict == 1) or (flag1dict == 1 and flag2geo ==1):
               flag1geo = 0
               flag1dict = 0
               flag2geo = 0
               flag2dict = 0
               continue
           for g in range(len(GeoDict)):
               if (splitTxt[i].find(GeoDict[g]) != -1):
                   flag1geo = 1
               if (splitTxt[i+1].find(GeoDict[g]) != -1):
                   flag2geo = 1
                    #List.append(splitTxt[i]+'.' + splitTxt[i + 1]+'.')
           for d in range(len(Dict)):
               if (splitTxt[i].find(Dict[d]) != -1):
                   flag1dict = 1
               if (splitTxt[i+1].find(Dict[d]) != -1):
                   flag2dict = 1
           if flag1geo == 1 and flag1dict == 1:
               List.append(splitTxt[i]+'.')
               flag1geo = 0
               flag1dict = 0
               flag2geo = 0
               flag2dict = 0
               continue
           elif((flag1geo == 1 and flag2dict == 1) or (flag1dict == 1 and flag2geo == 1)):
               List.append(splitTxt[i] + '.'+splitTxt[i+1]+'.')

        # for index in range(len(List)):
        #     for d in range(len(Dict)):
        #         res = List[index].find(Dict[d])
        #         if (res != -1):
        #             ListDict.append((List[index]))
        seen = set()
        for item in List:
            if item not in seen:
                seen.add(item)
                Result.append(item.replace(' e ','-').lstrip())

        for l in range(len(Result)):
            print(Result[l])
        print("ok")
        f = open('../middle/' + osfile[index].split('.')[0] + ".txt", "w", encoding='utf-8')
        for line in Result:
            f.write(line + '\n')
        print(osfile[index].split('.')[0] + ".txt" + "保存成功")
        f.close()
        print("okA")

# 展示文本信息
def showDocx(paragraphsText_new):
    for index, value in enumerate(paragraphsText_new):
        print(index, value)
# 保存文本数据
def saveTxt(paragraphsText_new,Namedocx):
    print(Namedocx)
    f = open('../last/'+Namedocx.split('.')[0]+".txt", "w", encoding='utf-8')
    for line in paragraphsText_new:
        f.write(line + '\n')
    print(Namedocx+"保存成功")
    f.close()

if __name__ == '__main__':
    osfile = []
    dir = r"D:\pythonProject\CarbonatePlatform\extracttext"
    for root, dirs, files in os.walk(dir):
        for file in files:
            if file != 'extract.py' and file != '__init__.py' and file !='split.py' and file !='Label.py' and file !='CSV.py' and file !='resplit.py' and file !='cluster.py':
                osfile.append(file)
    for index in range(len(osfile)):
        # 参数数组
        Dict = []
        GeoDict = []
        paragraphsText = []
        a = 0
        paragraphsText_new = []
        Result = []
        # 读取文本合并
        doc = Document(osfile[index])
        dict()
        geodict()
        getParagraphsText(doc,osfile[index])
        # 筛选
        filter()
        #展示存储
        showDocx(paragraphsText_new)
        saveTxt(paragraphsText_new, osfile[index])